"""
Document processing pipeline.

Handles document upload, text extraction, chunking, and embedding
for use as a verification source via pgvector.
"""

import logging
from typing import Optional
from dataclasses import dataclass, field

from app.config import get_settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """A single chunk of a processed document."""
    chunk_index: int
    content: str
    embedding: Optional[list[float]] = None


@dataclass 
class ProcessedDocument:
    """A fully processed document ready for storage."""
    filename: str
    file_type: str
    file_size_bytes: int
    chunks: list[DocumentChunk] = field(default_factory=list)


class DocumentProcessor:
    """
    Processes uploaded documents: extract text → chunk → embed.
    
    Supports PDF and DOCX files.
    """

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.settings = get_settings()

    async def process(
        self, content: bytes, filename: str, file_type: str
    ) -> ProcessedDocument:
        """
        Process an uploaded document.
        
        Args:
            content: Raw file bytes.
            filename: Original filename.
            file_type: MIME type or extension.
            
        Returns:
            ProcessedDocument with text chunks (embeddings added in Phase 2).
        """
        # Extract text
        text = self._extract_text(content, file_type)
        if not text:
            logger.warning(f"No text extracted from {filename}")
            return ProcessedDocument(
                filename=filename,
                file_type=file_type,
                file_size_bytes=len(content),
            )

        # Chunk the text
        chunks = self._chunk_text(text)

        doc = ProcessedDocument(
            filename=filename,
            file_type=file_type,
            file_size_bytes=len(content),
            chunks=chunks,
        )

        logger.info(f"Processed {filename}: {len(chunks)} chunks")
        return doc

    def _extract_text(self, content: bytes, file_type: str) -> str:
        """Extract text from document bytes."""
        file_type_lower = file_type.lower()

        if "pdf" in file_type_lower:
            return self._extract_pdf(content)
        elif "docx" in file_type_lower or "document" in file_type_lower:
            return self._extract_docx(content)
        elif "text" in file_type_lower or "txt" in file_type_lower:
            return content.decode("utf-8", errors="ignore")
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return content.decode("utf-8", errors="ignore")

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            from pypdf import PdfReader
            import io

            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(content))
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    def _chunk_text(self, text: str) -> list[DocumentChunk]:
        """Split text into overlapping chunks."""
        words = text.split()
        chunks = []
        idx = 0
        chunk_index = 0

        while idx < len(words):
            end = min(idx + self.chunk_size, len(words))
            chunk_text = " ".join(words[idx:end])

            chunks.append(DocumentChunk(
                chunk_index=chunk_index,
                content=chunk_text,
            ))

            chunk_index += 1
            idx += self.chunk_size - self.chunk_overlap

        return chunks


# ── Module-level singleton ────────────────────────────────────────────────

_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the document processor singleton."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
